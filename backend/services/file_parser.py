"""
file_parser.py — Extraction texte depuis pièces jointes (PDF / DOCX / TXT / MD).

Conventions :
  - Limite stricte par fichier : FILE_MAX_BYTES (par défaut 10 MiB).
  - Limite agrégée par requête : TOTAL_MAX_BYTES (par défaut 25 MiB).
  - Tronquage texte extrait : MAX_CHARS_PER_FILE (par défaut 30 000 caractères).
  - Aucune exécution de macro / script. python-docx et pypdf parsent uniquement le texte.
  - Le fichier brut n'est PAS persisté en base — seul le texte extrait l'est,
    et il transite chiffré dans laurentia_interactions (input_text).
"""
from __future__ import annotations

import io
import logging
from dataclasses import dataclass, asdict
from typing import Iterable

import pypdf
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

# Limites
FILE_MAX_BYTES = 10 * 1024 * 1024  # 10 MiB / fichier
TOTAL_MAX_BYTES = 25 * 1024 * 1024  # 25 MiB / requête
MAX_CHARS_PER_FILE = 30_000        # ~7.5k tokens

SUPPORTED_EXTENSIONS = {"pdf", "docx", "txt", "md", "markdown"}


class FileParseError(Exception):
    """Erreur d'extraction (format invalide, fichier corrompu, taille dépassée)."""


@dataclass
class ParsedFile:
    filename: str
    kind: str            # "pdf" | "docx" | "txt" | "md"
    bytes_size: int
    chars: int
    truncated: bool
    text: str

    def as_summary(self) -> dict:
        return {
            "filename": self.filename,
            "kind": self.kind,
            "bytes_size": self.bytes_size,
            "chars": self.chars,
            "truncated": self.truncated,
        }


def _guess_kind(filename: str, content_type: str | None) -> str:
    name = (filename or "").lower()
    ext = name.rsplit(".", 1)[-1] if "." in name else ""
    if ext in SUPPORTED_EXTENSIONS:
        return "md" if ext == "markdown" else ext
    if content_type:
        if "pdf" in content_type:
            return "pdf"
        if "officedocument.wordprocessingml" in content_type or content_type.endswith("/msword"):
            return "docx"
        if content_type.startswith("text/markdown"):
            return "md"
        if content_type.startswith("text/"):
            return "txt"
    raise FileParseError(f"Format non supporté : {filename}")


def _extract_pdf(data: bytes) -> str:
    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
    except Exception as e:
        raise FileParseError(f"PDF illisible : {e}") from e
    parts: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t:
            parts.append(t)
    return "\n\n".join(parts).strip()


def _extract_docx(data: bytes) -> str:
    try:
        doc = DocxDocument(io.BytesIO(data))
    except Exception as e:
        raise FileParseError(f"DOCX illisible : {e}") from e
    parts: list[str] = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _extract_text(data: bytes) -> str:
    try:
        return data.decode("utf-8", errors="replace").strip()
    except Exception as e:
        raise FileParseError(f"Texte illisible : {e}") from e


def parse_file(filename: str, content_type: str | None, data: bytes) -> ParsedFile:
    """Parse un fichier unique. Lève FileParseError sur format/limite invalide."""
    size = len(data)
    if size == 0:
        raise FileParseError(f"Fichier vide : {filename}")
    if size > FILE_MAX_BYTES:
        raise FileParseError(
            f"Fichier trop volumineux : {filename} ({size} > {FILE_MAX_BYTES} bytes)"
        )

    kind = _guess_kind(filename, content_type)
    if kind == "pdf":
        text = _extract_pdf(data)
    elif kind == "docx":
        text = _extract_docx(data)
    else:  # txt / md
        text = _extract_text(data)

    truncated = False
    if len(text) > MAX_CHARS_PER_FILE:
        text = text[:MAX_CHARS_PER_FILE]
        truncated = True

    return ParsedFile(
        filename=filename or "document",
        kind=kind,
        bytes_size=size,
        chars=len(text),
        truncated=truncated,
        text=text,
    )


def parse_many(files: Iterable[tuple[str, str | None, bytes]]) -> list[ParsedFile]:
    """
    Parse plusieurs fichiers, applique la limite agrégée TOTAL_MAX_BYTES.
    files : itérable de tuples (filename, content_type, raw_bytes)
    """
    parsed: list[ParsedFile] = []
    total = 0
    for filename, ctype, data in files:
        total += len(data)
        if total > TOTAL_MAX_BYTES:
            raise FileParseError(
                f"Volume total des pièces jointes trop important (> {TOTAL_MAX_BYTES} bytes)"
            )
        parsed.append(parse_file(filename, ctype, data))
    return parsed


def build_context_block(parsed: list[ParsedFile]) -> str:
    """
    Concatène le texte extrait dans un bloc Markdown lisible pour Claude.
    Forme :
        ## Pièces jointes
        ### document1.pdf (pdf, 8 920 chars[, tronqué])
        <contenu>
        ---
        ### note.md (md, 312 chars)
        <contenu>
    """
    if not parsed:
        return ""
    lines: list[str] = ["## Pièces jointes utilisateur"]
    for pf in parsed:
        flag = ", tronqué" if pf.truncated else ""
        lines.append(f"### {pf.filename} ({pf.kind}, {pf.chars} chars{flag})")
        lines.append(pf.text)
        lines.append("\n---\n")
    return "\n".join(lines).strip()
