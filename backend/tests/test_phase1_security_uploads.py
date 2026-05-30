"""
Tests Phase 1 — Chiffrement AES-256-GCM + Upload de fichiers (PDF/DOCX/TXT).

Couvre :
  - crypto.encrypt_text / decrypt_text round-trip + rétro-compat plaintext
  - parser PDF/DOCX/TXT/MD + erreurs (taille, format)
  - Endpoint /api/laurentia/upload-limits public
  - Endpoint /api/laurentia/query multipart : gate tier + extraction visible
    pour Claude (via assertion sur le SSE).
  - Persistence : le champ input_text en DB doit être un dict {v:1,n,c} (chiffré).
"""
import io
import json
import os

import httpx
import pytest
from pymongo import MongoClient

from services.crypto import encrypt_text, decrypt_text, is_encrypted
from services.file_parser import (
    parse_file,
    parse_many,
    FileParseError,
    build_context_block,
    FILE_MAX_BYTES,
)

BASE_URL = os.environ.get("REACT_APP_BACKEND_URL", "http://localhost:8001").rstrip("/")
MONGO_URL = os.environ.get("MONGO_URL", "mongodb://localhost:27017")
DB_NAME = os.environ.get("DB_NAME", "laurentia")

DEMO_FREK = "DEMO-SAYD"


# -------------------- Crypto --------------------

def test_encrypt_decrypt_round_trip():
    blob = encrypt_text("hello souverain")
    assert is_encrypted(blob)
    assert blob["v"] == 1
    assert "n" in blob and "c" in blob
    assert decrypt_text(blob) == "hello souverain"


def test_encrypt_text_none():
    assert encrypt_text(None) is None
    assert decrypt_text(None) == ""


def test_decrypt_legacy_plaintext():
    """Rétro-compatibilité : str legacy renvoyé tel quel."""
    assert decrypt_text("legacy plaintext message") == "legacy plaintext message"


def test_encrypt_each_call_uses_unique_nonce():
    a = encrypt_text("same input")
    b = encrypt_text("same input")
    assert a["n"] != b["n"]
    assert a["c"] != b["c"]
    assert decrypt_text(a) == decrypt_text(b) == "same input"


# -------------------- File parser --------------------

def test_parse_text_file():
    pf = parse_file("notes.txt", "text/plain", b"ligne 1\nligne 2")
    assert pf.kind == "txt"
    assert "ligne 1" in pf.text
    assert pf.chars > 0
    assert pf.truncated is False


def test_parse_md_file():
    pf = parse_file("note.md", "text/markdown", b"# Titre\nContenu")
    assert pf.kind == "md"
    assert "# Titre" in pf.text


def test_parse_empty_file_rejected():
    with pytest.raises(FileParseError):
        parse_file("empty.txt", "text/plain", b"")


def test_parse_unsupported_format():
    with pytest.raises(FileParseError):
        parse_file("photo.png", "image/png", b"\x89PNG\r\n")


def test_parse_oversize_file():
    big = b"x" * (FILE_MAX_BYTES + 1)
    with pytest.raises(FileParseError):
        parse_file("big.txt", "text/plain", big)


def test_parse_pdf():
    """Génère un PDF minimaliste avec reportlab, le parse, vérifie le texte."""
    try:
        from reportlab.pdfgen import canvas
    except ImportError:
        pytest.skip("reportlab non installé pour ce test")
    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.drawString(100, 750, "Doctrine Laurent.ia v0.8.")
    c.drawString(100, 730, "Diaspora = 42 pourcents PIB.")
    c.save()
    data = buf.getvalue()
    pf = parse_file("doctrine.pdf", "application/pdf", data)
    assert pf.kind == "pdf"
    assert "Doctrine Laurent.ia" in pf.text
    assert "42 pourcents" in pf.text


def test_parse_docx():
    """Génère un DOCX, le parse."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        pytest.skip("python-docx non installé pour ce test")
    doc = DocxDocument()
    doc.add_heading("Brief", level=1)
    doc.add_paragraph("Phrase test : caraïbe souveraine.")
    buf = io.BytesIO()
    doc.save(buf)
    pf = parse_file(
        "brief.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        buf.getvalue(),
    )
    assert pf.kind == "docx"
    assert "caraïbe souveraine" in pf.text


def test_build_context_block_format():
    pf = parse_file("a.txt", "text/plain", b"alpha")
    block = build_context_block([pf])
    assert "## Pièces jointes utilisateur" in block
    assert "### a.txt" in block
    assert "alpha" in block


def test_parse_many_aggregate_limit():
    """Limite agrégée doit lever quand somme > TOTAL_MAX_BYTES."""
    big = b"x" * (FILE_MAX_BYTES)  # juste sous la limite par fichier
    # On envoie 3 fichiers de FILE_MAX_BYTES → 30 Mio > TOTAL_MAX_BYTES (25 Mio)
    files = [(f"f{i}.txt", "text/plain", big) for i in range(3)]
    with pytest.raises(FileParseError):
        parse_many(files)


# -------------------- Endpoints HTTP --------------------

@pytest.fixture(scope="module")
def http_client():
    return httpx.Client(base_url=BASE_URL, timeout=30.0)


@pytest.fixture(scope="module")
def db():
    client = MongoClient(MONGO_URL)
    yield client[DB_NAME]
    client.close()


def test_upload_limits_endpoint(http_client):
    r = http_client.get("/api/laurentia/upload-limits")
    assert r.status_code == 200
    body = r.json()
    assert "file_max_bytes" in body and body["file_max_bytes"] > 0
    assert "allowed_tiers" in body and set(body["allowed_tiers"]) == {"creator", "infinite"}
    assert "pdf" in body["extensions"]


def _set_tier(db, frek_id, tier):
    quotas = {
        "free": dict(tier="free", version="free", tokens_limit_month=100_000, memory_window=10, rate_per_min=10),
        "creator": dict(tier="creator", version="creator", tokens_limit_month=2_000_000, memory_window=100, rate_per_min=30),
    }
    db.laurentia_instances.update_one(
        {"frek_id": frek_id},
        {"$set": quotas[tier], "$setOnInsert": {"frek_id": frek_id}},
        upsert=True,
    )


def test_query_multipart_free_tier_rejected(http_client, db):
    _set_tier(db, DEMO_FREK, "free")
    payload = json.dumps({"frek_id": DEMO_FREK, "input": "test", "context": {"app": "direct"}})
    files = [("files", ("note.txt", b"x", "text/plain"))]
    r = http_client.post("/api/laurentia/query", data={"payload": payload}, files=files)
    assert r.status_code == 403
    assert "Creator" in r.text or "creator" in r.text


def test_query_multipart_creator_ok_and_encrypted_at_rest(http_client, db):
    _set_tier(db, DEMO_FREK, "creator")
    payload = json.dumps({"frek_id": DEMO_FREK, "input": "Cite la phrase exacte", "context": {"app": "direct"}})
    files = [("files", ("doc.txt", b"Mot-clef secret: AZALEA-2026.", "text/plain"))]
    r = http_client.post("/api/laurentia/query", data={"payload": payload}, files=files)
    assert r.status_code == 200
    body = r.text
    # SSE meta annonce le fichier
    assert "event: meta" in body
    assert "doc.txt" in body
    assert "event: done" in body

    # Vérifie l'enregistrement chiffré
    rec = db.laurentia_interactions.find_one(
        {"context_app": "direct"}, sort=[("timestamp", -1)]
    )
    assert rec is not None
    assert isinstance(rec.get("input_text"), dict)
    assert rec["input_text"].get("v") == 1
    assert "n" in rec["input_text"] and "c" in rec["input_text"]
    # Le texte clair (incluant le contenu du fichier) doit pouvoir se déchiffrer
    plain = decrypt_text(rec["input_text"])
    assert "Cite la phrase exacte" in plain
    assert "AZALEA-2026" in plain  # le bloc fichiers a bien été intégré dans input_text


def test_query_json_path_still_works(http_client, db):
    _set_tier(db, DEMO_FREK, "free")
    r = http_client.post(
        "/api/laurentia/query",
        json={"frek_id": DEMO_FREK, "input": "Réponds OK", "context": {"app": "direct"}},
    )
    assert r.status_code == 200
    assert "event: meta" in r.text
    assert "event: done" in r.text
